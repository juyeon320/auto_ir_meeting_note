#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
IR 미팅 녹취록(클로바노트 등에서 추출한 txt) -> 회의록 docx + 이메일 초안 자동 생성

사용법:
    export ANTHROPIC_API_KEY="sk-ant-..."
    python process_meeting.py transcript.txt --config config.json --outdir ./output

필요 패키지:
    pip install anthropic python-docx
"""

import argparse
import json
import re
import sys
from pathlib import Path

from anthropic import Anthropic
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from prompts import (
    QA_EXTRACTION_SYSTEM_PROMPT,
    TOPIC_SUMMARY_SYSTEM_PROMPT,
    build_user_message,
)

MODEL = "claude-sonnet-4-5"  # 필요시 최신 모델 스트링으로 교체하세요


def call_claude(client, system_prompt, user_message, max_tokens=8000):
    resp = client.messages.create(
        model=MODEL,
        max_tokens=max_tokens,
        system=system_prompt,
        messages=[{"role": "user", "content": user_message}],
    )
    text = "".join(block.text for block in resp.content if block.type == "text")
    return text


def extract_json(text: str):
    """모델 출력에서 JSON 부분만 안전하게 추출."""
    text = text.strip()
    text = re.sub(r"^```(json)?", "", text).strip()
    text = re.sub(r"```$", "", text).strip()
    return json.loads(text)


def extract_qa(client, transcript: str):
    user_msg = build_user_message(transcript)
    raw = call_claude(client, QA_EXTRACTION_SYSTEM_PROMPT, user_msg)
    return extract_json(raw)


def summarize_topics(client, qa_list):
    qa_text = "\n".join(f"Q: {item['q']}\nA: {item['a']}" for item in qa_list)
    raw = call_claude(client, TOPIC_SUMMARY_SYSTEM_PROMPT, qa_text, max_tokens=1000)
    data = extract_json(raw)
    return data["bullets"]


def _set_korean_font(doc, font_name="맑은 고딕"):
    """기본 스타일 폰트를 한글 폰트로 지정 (중간점 · 등 특수문자 렌더링 문제 방지)"""
    style = doc.styles["Normal"]
    style.font.name = font_name
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


def build_docx(meta: dict, qa_list: list, topic_bullets: list, output_path: Path):
    doc = Document()
    _set_korean_font(doc)

    title = doc.add_paragraph()
    run = title.add_run("IR 회의록")
    run.bold = True
    run.font.size = Pt(16)

    def add_line(label, value):
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(value)

    doc.add_paragraph()  # spacer
    add_line("1.일시", meta["datetime"])
    add_line("2.IR형태", meta.get("ir_type", "기업탐방"))
    add_line("3.IR대상기관", meta["target_org"])
    add_line("4.참석인원", meta["attendees"])
    add_line("5.특이사항", meta.get("notes", "없음"))

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("6. 주요 논의 내용").bold = True
    for i, bullet in enumerate(topic_bullets, 1):
        doc.add_paragraph(f"  {i}) {bullet}")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("7. Q&A 요약").bold = True
    doc.add_paragraph()

    for i, item in enumerate(qa_list, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(14)

        r_qlabel = p.add_run(f"Q{i}.")
        r_qlabel.bold = True
        r_qlabel.add_break()  # 줄바꿈: Q라벨 -> 질문

        r_qtext = p.add_run(item["q"])
        r_qtext.add_break()  # 줄바꿈: 질문 -> A라벨

        r_alabel = p.add_run(f"A{i}.")
        r_alabel.bold = True
        r_alabel.add_break()  # 줄바꿈: A라벨 -> 답변

        p.add_run(item["a"])
        # 이 paragraph가 끝나면 space_after(14pt)로 다음 Q/A 세트와 간격이 생김

    doc.save(output_path)


def build_email_draft(meta: dict, output_path: Path):
    email = f"""제목: {meta['email_date_short']} {meta['target_org']} IR미팅 회의록 송부드립니다.

수신자 제위.
안녕하세요 {meta['sender_name']}입니다.

{meta['email_date_long']} {meta['target_org']} 기업탐방 IR 미팅 내용을 송부드립니다.

1. 일시: {meta['datetime']}
2. 참석인원: {meta['attendees']}
3. 주요내용:
{meta['email_main_content']}
4. 후속조치: 특이사항 없음
5. 기타: 특이사항 없음
본 메일에 파일 첨부하였습니다.
첨부파일:

* {meta['attachment_base']} (docx)
* {meta['attachment_base']} (pdf)

답변 내용중 추가적으로 궁금하신 사항이 있으시면 회신 부탁드립니다.
감사합니다.
{meta['sender_name']} 드림.
"""
    output_path.write_text(email, encoding="utf-8")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("transcript", help="클로바노트 등에서 추출한 txt 파일 경로")
    parser.add_argument("--config", required=True, help="회의 메타정보 JSON 파일 경로")
    parser.add_argument("--outdir", default="./output", help="출력 디렉토리")
    args = parser.parse_args()

    transcript = Path(args.transcript).read_text(encoding="utf-8")
    meta = json.loads(Path(args.config).read_text(encoding="utf-8"))

    outdir = Path(args.outdir)
    outdir.mkdir(parents=True, exist_ok=True)

    client = Anthropic()  # ANTHROPIC_API_KEY 환경변수 사용

    print("[1/3] Q&A 추출 중...")
    qa_list = extract_qa(client, transcript)
    print(f"  -> {len(qa_list)}개 Q&A 추출 완료")

    print("[2/3] 주요 논의 내용 요약 중...")
    topic_bullets = summarize_topics(client, qa_list)

    # config에 email_main_content를 직접 안 넣었으면 topic_bullets로 자동 생성
    if not meta.get("email_main_content"):
        meta["email_main_content"] = "\n".join(
            f"  {i}) {b}" for i, b in enumerate(topic_bullets, 1)
        )

    print("[3/3] 문서 생성 중...")
    docx_path = outdir / f"{meta['file_base']}.docx"
    build_docx(meta, qa_list, topic_bullets, docx_path)

    email_path = outdir / f"{meta['file_base']}_이메일초안.txt"
    build_email_draft(meta, email_path)

    # 참고용 원본 JSON도 저장 (검수/재사용 편의)
    (outdir / f"{meta['file_base']}_qa.json").write_text(
        json.dumps(qa_list, ensure_ascii=False, indent=2), encoding="utf-8"
    )

    print(f"완료! 결과물: {outdir}")


if __name__ == "__main__":
    main()
