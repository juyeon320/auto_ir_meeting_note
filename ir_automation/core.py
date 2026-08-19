# -*- coding: utf-8 -*-
"""
IR 미팅 자동 정리 - 공용 핵심 로직.
CLI(process_meeting.py)와 웹앱(webapp/app.py)이 이 모듈을 공유합니다.
"""

import json
import re
from pathlib import Path

from openai import OpenAI
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from prompts import (
    QA_EXTRACTION_SYSTEM_PROMPT,
    TOPIC_SUMMARY_SYSTEM_PROMPT,
    build_user_message,
)

MODEL = "gpt-5-mini"  # 필요시 원하는 모델로 교체하세요 (gpt-5, gpt-5-mini, gpt-5-nano 등)


# ---------------------------------------------------------------------------
# OpenAI 호출
# ---------------------------------------------------------------------------

def call_openai(client, system_prompt, user_message, max_tokens=8000, force_json=False):
    kwargs = {}
    if force_json:
        kwargs["response_format"] = {"type": "json_object"}

    resp = client.chat.completions.create(
        model=MODEL,
        max_completion_tokens=max_tokens,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_message},
        ],
        **kwargs,
    )
    return resp.choices[0].message.content


def extract_json(text: str):
    text = text.strip()
    text = re.sub(r"^```(json)?", "", text).strip()
    text = re.sub(r"```$", "", text).strip()
    return json.loads(text)


def extract_qa(client, transcript: str):
    user_msg = build_user_message(transcript)
    raw = call_openai(client, QA_EXTRACTION_SYSTEM_PROMPT, user_msg, force_json=True)
    data = extract_json(raw)
    # 모델이 {"qa": [...]} 형태로 감싸서 줄 수도 있어 양쪽 다 처리
    return data["qa"] if isinstance(data, dict) and "qa" in data else data


def summarize_topics(client, qa_list):
    qa_text = "\n".join(f"Q: {item['q']}\nA: {item['a']}" for item in qa_list)
    raw = call_openai(client, TOPIC_SUMMARY_SYSTEM_PROMPT, qa_text, max_tokens=1000, force_json=True)
    return extract_json(raw)["bullets"]


def run_pipeline(client, transcript: str):
    """transcript -> (qa_list, topic_bullets). 웹앱/CLI 둘 다 이걸 호출하면 됨."""
    qa_list = extract_qa(client, transcript)
    topic_bullets = summarize_topics(client, qa_list)
    return qa_list, topic_bullets


# ---------------------------------------------------------------------------
# 문서 생성
# ---------------------------------------------------------------------------

def _set_korean_font(doc, font_name="맑은 고딕"):
    """기본 스타일 폰트를 한글 폰트로 지정 (중간점 · 등 특수문자 렌더링 문제 방지)"""
    style = doc.styles["Normal"]
    style.font.name = font_name
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


def build_docx(meta: dict, qa_list: list, topic_bullets: list, output_path: Path):
    doc = Document()
    _set_korean_font(doc)

    title = doc.add_paragraph()
    run = title.add_run("IR 회의록")
    run.bold = True
    run.font.size = Pt(16)

    def add_line(label, value):
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(value)

    doc.add_paragraph()
    add_line("1.일시", meta["datetime"])
    add_line("2.IR형태", meta.get("ir_type", "기업탐방"))
    add_line("3.IR대상기관", meta["target_org"])
    add_line("4.참석인원", meta["attendees"])
    add_line("5.특이사항", meta.get("notes", "없음"))

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("6. 주요 논의 내용").bold = True
    for i, bullet in enumerate(topic_bullets, 1):
        doc.add_paragraph(f"  {i}) {bullet}")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("7. Q&A 요약").bold = True
    doc.add_paragraph()

    for i, item in enumerate(qa_list, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(14)

        r_qlabel = p.add_run(f"Q{i}.")
        r_qlabel.bold = True
        r_qlabel.add_break()

        r_qtext = p.add_run(item["q"])
        r_qtext.add_break()

        r_alabel = p.add_run(f"A{i}.")
        r_alabel.bold = True
        r_alabel.add_break()

        p.add_run(item["a"])

    doc.save(output_path)


def build_email_draft(meta: dict, output_path: Path):
    email = f"""제목: {meta['email_date_short']} {meta['target_org']} IR미팅 회의록 송부드립니다.

수신자 제위.
안녕하세요 {meta['sender_name']}입니다.

{meta['email_date_long']} {meta['target_org']} 기업탐방 IR 미팅 내용을 송부드립니다.

1. 일시: {meta['datetime']}
2. 참석인원: {meta['attendees']}
3. 주요내용:
{meta['email_main_content']}
4. 후속조치: 특이사항 없음
5. 기타: 특이사항 없음
본 메일에 파일 첨부하였습니다.
첨부파일:

* {meta['attachment_base']} (docx)
* {meta['attachment_base']} (pdf)

답변 내용중 추가적으로 궁금하신 사항이 있으시면 회신 부탁드립니다.
감사합니다.
{meta['sender_name']} 드림.
"""
    output_path.write_text(email, encoding="utf-8")


def fill_email_main_content(meta: dict, topic_bullets: list):
    """config에 email_main_content가 없으면 topic_bullets로 자동 채움."""
    if not meta.get("email_main_content"):
        meta["email_main_content"] = "\n".join(
            f"  {i}) {b}" for i, b in enumerate(topic_bullets, 1)
        )
    return meta
